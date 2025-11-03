import customtkinter as ctk
from tkinter import filedialog, messagebox, simpledialog, PhotoImage
import re
from docx import Document
from fpdf import FPDF
import os
import copy

class WordEditorMassive(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Editor Word Masivo")
        self.geometry("900x600")

        # ICONO.......................
        try:
            icon_path = r"..\assets\icons\icon.png"
            if not os.path.exists(icon_path):
                print(f"⚠️ No se encontró el archivo del icono en: {icon_path}")
            else:
                icon_image = PhotoImage(file=icon_path)
                self.iconphoto(True, icon_image)
                print("✅ Icono cargado correctamente.")
        except Exception as e:
            print(f"❌ Error al cargar el icono: {e}")

        # Variables..................
        self.docs = []
        self.file_paths = []
        self.current_index = 0
        self.history = []  # Para deshacer cambios

        # Botones de carga arriba
        top_frame = ctk.CTkFrame(self)
        top_frame.pack(padx=10, pady=5, fill="x")
        self.load_btn = ctk.CTkButton(top_frame, text="Cargar Word", command=self.load_word)
        self.load_btn.pack(side="left", padx=5)
        self.load_many_btn = ctk.CTkButton(top_frame, text="Cargar Varios Word", command=self.load_many_words)
        self.load_many_btn.pack(side="left", padx=5)

        # Entry de búsqueda y reemplazo
        self.search_entry = ctk.CTkEntry(self, placeholder_text="Buscar...")
        self.search_entry.pack(padx=10, pady=5, fill="x")
        self.search_entry.bind("<KeyRelease>", lambda e: self.highlight_matches())
        self.replace_entry = ctk.CTkEntry(self, placeholder_text="Reemplazar por...")
        self.replace_entry.pack(padx=10, pady=5, fill="x")

        # Botones horizontales
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(padx=10, pady=5, fill="x")
        self.replace_btn = ctk.CTkButton(btn_frame, text="Reemplazar", command=self.replace_text_selective)
        self.replace_btn.pack(side="left", padx=5)
        self.replace_all_btn = ctk.CTkButton(btn_frame, text="Reemplazar Todo", command=self.replace_all_docs)
        self.replace_all_btn.pack(side="left", padx=5)
        self.undo_btn = ctk.CTkButton(btn_frame, text="Deshacer", command=self.undo)
        self.undo_btn.pack(side="left", padx=5)
        self.clear_btn = ctk.CTkButton(btn_frame, text="Limpiar", command=self.clear_text)
        self.clear_btn.pack(side="left", padx=5)
        self.save_btn = ctk.CTkButton(btn_frame, text="Guardar Como", command=self.save_file)
        self.save_btn.pack(side="left", padx=5)

        # Navegación entre archivos
        nav_frame = ctk.CTkFrame(self)
        nav_frame.pack(padx=10, pady=5, fill="x")
        self.prev_btn = ctk.CTkButton(nav_frame, text="<< Anterior", command=self.prev_doc)
        self.prev_btn.pack(side="left", padx=5)
        self.next_btn = ctk.CTkButton(nav_frame, text="Siguiente >>", command=self.next_doc)
        self.next_btn.pack(side="left", padx=5)
        self.file_label = ctk.CTkLabel(nav_frame, text="No hay archivos cargados")
        self.file_label.pack(side="left", padx=10)

        # TextArea
        self.text_area = ctk.CTkTextbox(self, wrap="word")
        self.text_area.pack(padx=10, pady=10, fill="both", expand=True)

    #  Funciones .................
    def load_word(self):
        path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if path:
            self.file_paths = [path]
            self.docs = [Document(path)]
            self.current_index = 0
            self.history = [copy.deepcopy(self.docs)]
            self.update_text_area()
            self.update_file_label()

    def load_many_words(self):
        paths = filedialog.askopenfilenames(filetypes=[("Word Files", "*.docx")])
        if paths:
            self.file_paths = list(paths)
            self.docs = [Document(p) for p in paths]
            self.current_index = 0
            self.history = [copy.deepcopy(self.docs)]
            self.update_text_area()
            self.update_file_label()

    def update_text_area(self):
        if not self.docs:
            return
        self.text_area.delete("1.0", "end")
        for para in self.docs[self.current_index].paragraphs:
            self.text_area.insert("end", para.text + "\n\n")
        self.highlight_matches()

    def update_file_label(self):
        if self.file_paths:
            self.file_label.configure(
                text=f"Archivo {self.current_index +1}/{len(self.file_paths)}: {os.path.basename(self.file_paths[self.current_index])}"
            )
        else:
            self.file_label.configure(text="No hay archivos cargados")

    def highlight_matches(self):
        self.text_area.tag_remove("highlight", "1.0", "end")
        search_text = self.search_entry.get()
        if not search_text:
            return
        start_pos = "1.0"
        pattern = re.compile(re.escape(search_text), re.IGNORECASE)
        while True:
            match = pattern.search(self.text_area.get(start_pos, "end"))
            if not match:
                break
            start_index = f"{start_pos}+{match.start()}c"
            end_index = f"{start_pos}+{match.end()}c"
            self.text_area.tag_add("highlight", start_index, end_index)
            start_pos = end_index
        self.text_area.tag_config("highlight", background="yellow")

    def replace_text_selective(self):
        if not self.docs:
            messagebox.showwarning("Aviso", "Carga primero un archivo Word")
            return
        search_text = self.search_entry.get()
        replace_text = self.replace_entry.get()
        if not search_text:
            messagebox.showinfo("Info", "Ingresa texto a buscar")
            return
        doc = self.docs[self.current_index]
        pattern = re.compile(re.escape(search_text), re.IGNORECASE)
        matches = []
        self.history.append(copy.deepcopy(self.docs))
        for i, para in enumerate(doc.paragraphs):
            for m in pattern.finditer(para.text):
                matches.append((i, m.start(), m.end(), para.text[m.start():m.end()]))
        if not matches:
            messagebox.showinfo("Info", "No se encontraron coincidencias")
            return
        for idx, (p_idx, start, end, text) in enumerate(matches, start=1):
            answer = messagebox.askyesno("Reemplazar coincidencia", f"Reemplazar '{text}' en el párrafo {p_idx+1}?")
            if answer:
                para = doc.paragraphs[p_idx]
                para.text = para.text[:start] + replace_text + para.text[end:]
        self.update_text_area()
        messagebox.showinfo("Reemplazo", "Reemplazo selectivo completado.")

    def replace_all_docs(self):
        if not self.docs:
            messagebox.showwarning("Aviso", "Carga primero uno o varios archivos Word")
            return
        search_text = self.search_entry.get()
        replace_text = self.replace_entry.get()
        if not search_text:
            messagebox.showinfo("Info", "Ingresa texto a buscar")
            return
        pattern = re.compile(re.escape(search_text), re.IGNORECASE)
        total_replaced = 0
        self.history.append(copy.deepcopy(self.docs))
        for doc in self.docs:
            for para in doc.paragraphs:
                if pattern.search(para.text):
                    para.text = pattern.sub(replace_text, para.text)
                    total_replaced += 1
        self.update_text_area()
        messagebox.showinfo("Reemplazo Masivo", f"Se reemplazaron {total_replaced} coincidencias en todos los documentos.")

    def undo(self):
        if len(self.history) > 1:
            self.history.pop()
            self.docs = copy.deepcopy(self.history[-1])
            self.update_text_area()
            messagebox.showinfo("Deshacer", "Cambio deshecho.")
        else:
            messagebox.showinfo("Deshacer", "No hay cambios para deshacer.")

    def clear_text(self):
        self.text_area.delete("1.0", "end")
        self.search_entry.delete(0, "end")
        self.replace_entry.delete(0, "end")
        self.text_area.tag_remove("highlight", "1.0", "end")

    def save_file(self):
        if not self.docs:
            messagebox.showwarning("Aviso", "Carga primero un archivo Word")
            return
        save_dir = filedialog.askdirectory(title="Selecciona carpeta para guardar")
        if save_dir:
            for i, doc in enumerate(self.docs):
                filename = os.path.basename(self.file_paths[i])
                save_path_docx = os.path.join(save_dir, filename)
                doc.save(save_path_docx)
            messagebox.showinfo("Guardado", f"Archivos guardados en {save_dir}")

    def prev_doc(self):
        if self.current_index > 0:
            self.current_index -= 1
            self.update_text_area()
            self.update_file_label()

    def next_doc(self):
        if self.current_index < len(self.docs) - 1:
            self.current_index += 1
            self.update_text_area()
            self.update_file_label()


if __name__ == "__main__":
    app = WordEditorMassive()
    app.mainloop()
